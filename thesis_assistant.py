from __future__ import annotations

import argparse
import csv
import hashlib
import json
import math
import re
import shutil
import sys
from collections import Counter
from dataclasses import asdict, dataclass
from datetime import datetime, timezone
from difflib import SequenceMatcher
from pathlib import Path
from typing import Iterable


ROOT = Path(__file__).resolve().parent
PAPERS = ROOT / "papers"
DRAFTS = ROOT / "drafts"
OUTPUT = ROOT / "output"
INDEX_FILE = OUTPUT / "library_index.json"
AUDIT_FILE = OUTPUT / "citation_audit.csv"
WEB_QUEUE = OUTPUT / "web_search_requests.jsonl"
CITATION_NEED_THRESHOLD = 25


def configure_paths(papers_dir: Path | None = None, output_dir: Path | None = None) -> None:
    """Configure a reusable paper library and a separate report directory."""
    global PAPERS, OUTPUT, INDEX_FILE, AUDIT_FILE, WEB_QUEUE
    if papers_dir is not None:
        PAPERS = papers_dir.expanduser().resolve()
    if output_dir is not None:
        OUTPUT = output_dir.expanduser().resolve()
    INDEX_FILE = OUTPUT / "library_index.json"
    AUDIT_FILE = OUTPUT / "citation_audit.csv"
    WEB_QUEUE = OUTPUT / "web_search_requests.jsonl"


CLAIM_RULES = {
    "quantitative": re.compile(r"(?:\d[\d,]*(?:\.\d+)?\s*(?:%|nm|μm|um|mm|cm|eV|V|mA|A|K|°C|PPI|EQE|cd/?m2|mol%))", re.I),
    "comparison": re.compile(r"(?:보다|대비|더\s|높(?:다|은|였다)|낮(?:다|은|았다)|우수|향상|감소|증가|개선|효율적|간단|compared|higher|lower|superior|improv|increase|decrease|enhanc|limit(?:ed|ation)|drawback|promising|strong candidate|high-resolution|high-efficiency)", re.I),
    "causal": re.compile(r"(?:때문에|따라서|유도|기인|영향을\s*주|증가시|감소시|결과로|로\s*인해|owing to|due to|therefore|thereby|leads? to|results? in|enabl(?:e|es|ed)|necessitat)", re.I),
    "generalization": re.compile(r"(?:널리|일반적으로|주로|대부분|핵심\s*기술|필수적|알려져|보고되|typically|generally|widely|commonly|has emerged|have emerged|demonstrat(?:e|es|ed)|reported|require(?:s|d)?|demand(?:s|ed)?)", re.I),
    "mechanism": re.compile(r"(?:메커니즘|mechanism|가교|crosslink|ligand|리간드|결함|전하|전자|정공|여기자|반응)", re.I),
}

VERDICT_LABELS = {
    "direct": "직접 근거",
    "partial": "부분 근거",
    "mismatch": "근거 불일치",
    "insufficient": "판단 자료 부족",
}

CONFIDENCE_LABELS = {"high": "높음", "medium": "보통", "low": "낮음"}

ACTION_LABELS = {
    "keep": "유지",
    "revise": "문장 수정",
    "split": "문장 분리",
    "add_reference": "근거 추가",
    "replace_reference": "근거 교체",
    "external_search": "외부 논문 검색",
    "review_external_proposals": "외부 논문 제안 검토",
}

CATEGORY_LABELS = {
    "quantitative": "정량 주장",
    "comparison": "비교 주장",
    "causal": "인과 주장",
    "generalization": "일반화 주장",
    "mechanism": "메커니즘 주장",
    "existing_citation": "기존 인용 있음",
}


def display_label(value: str, labels: dict[str, str]) -> str:
    return labels.get(value, value or "미정")


@dataclass
class Chunk:
    id: str
    file: str
    page: int
    text: str
    title: str = ""
    doi: str = ""
    year: str = ""


@dataclass
class Hit:
    chunk: Chunk
    score: float


def ensure_dirs() -> None:
    for path in (PAPERS, DRAFTS, OUTPUT):
        path.mkdir(parents=True, exist_ok=True)


def normalize_space(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def tokenize(text: str) -> list[str]:
    lowered = text.lower()
    latin = re.findall(r"[a-z][a-z0-9+.-]{1,}|\d+(?:\.\d+)?|μm|°c", lowered)
    hangul_runs = re.findall(r"[가-힣]{2,}", lowered)
    grams: list[str] = []
    for run in hangul_runs:
        for size in (2, 3):
            grams.extend(run[i : i + size] for i in range(len(run) - size + 1))
    return latin + grams


def split_chunks(page_text: str, max_chars: int = 1400) -> list[str]:
    paragraphs = [normalize_space(p) for p in re.split(r"\n\s*\n|(?<=[.!?])\s+(?=[A-Z가-힣])", page_text)]
    # Figure captions and concise quantitative findings can be short but valuable.
    paragraphs = [p for p in paragraphs if len(p) >= 15]
    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs:
        if current and len(current) + len(paragraph) + 1 > max_chars:
            chunks.append(current)
            current = paragraph
        else:
            current = f"{current} {paragraph}".strip()
    if current:
        chunks.append(current)
    return chunks


def extract_metadata(first_pages: str, filename: str) -> tuple[str, str, str]:
    doi_match = re.search(r"\b10\.\d{4,9}/[-._;()/:A-Z0-9]+", first_pages, re.I)
    year_match = re.search(r"\b(?:19|20)\d{2}\b", first_pages)
    lines = [normalize_space(line) for line in first_pages.splitlines() if len(normalize_space(line)) > 10]
    title = next((line for line in lines[:20] if 25 <= len(line) <= 250 and not line.lower().startswith(("abstract", "doi"))), Path(filename).stem)
    return title, doi_match.group(0).rstrip(".,;") if doi_match else "", year_match.group(0) if year_match else ""


def index_papers() -> list[Chunk]:
    ensure_dirs()
    try:
        import pdfplumber
    except ImportError as exc:
        raise SystemExit("pdfplumber가 필요함: python -m pip install pdfplumber") from exc

    pdfs = sorted(PAPERS.glob("*.pdf"))
    if not pdfs:
        raise SystemExit(f"PDF가 없음: {PAPERS}")

    indexed: list[Chunk] = []
    errors: list[str] = []
    for pdf in pdfs:
        try:
            with pdfplumber.open(pdf) as doc:
                page_texts = [(page.extract_text() or "") for page in doc.pages]
            title, doi, year = extract_metadata("\n".join(page_texts[:2]), pdf.name)
            for page_number, page_text in enumerate(page_texts, start=1):
                for part, text in enumerate(split_chunks(page_text), start=1):
                    digest = hashlib.sha1(f"{pdf.name}:{page_number}:{part}".encode("utf-8")).hexdigest()[:12]
                    indexed.append(Chunk(digest, pdf.name, page_number, text, title, doi, year))
        except Exception as exc:  # keep other valid papers usable
            errors.append(f"{pdf.name}: {exc}")

    payload = {
        "created_at": datetime.now(timezone.utc).isoformat(),
        "paper_count": len(pdfs),
        "chunk_count": len(indexed),
        "errors": errors,
        "chunks": [asdict(chunk) for chunk in indexed],
    }
    INDEX_FILE.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"색인 완료: PDF {len(pdfs)}개, 텍스트 조각 {len(indexed)}개")
    if errors:
        print("추출 실패:")
        for error in errors:
            print(f"- {error}")
    return indexed


def load_index() -> list[Chunk]:
    if not INDEX_FILE.exists():
        return index_papers()
    data = json.loads(INDEX_FILE.read_text(encoding="utf-8"))
    return [Chunk(**item) for item in data["chunks"]]


class BM25:
    def __init__(self, chunks: list[Chunk]):
        self.chunks = chunks
        self.docs = [tokenize(chunk.text) for chunk in chunks]
        self.counts = [Counter(doc) for doc in self.docs]
        self.avgdl = sum(map(len, self.docs)) / max(len(self.docs), 1)
        document_frequency: Counter[str] = Counter()
        for doc in self.docs:
            document_frequency.update(set(doc))
        size = len(self.docs)
        self.idf = {term: math.log(1 + (size - freq + 0.5) / (freq + 0.5)) for term, freq in document_frequency.items()}

    def search(self, query: str, limit: int = 3) -> list[Hit]:
        query_terms = set(tokenize(query))
        hits: list[Hit] = []
        for chunk, counts, doc in zip(self.chunks, self.counts, self.docs):
            score = 0.0
            doc_len = len(doc)
            for term in query_terms:
                freq = counts.get(term, 0)
                if not freq:
                    continue
                score += self.idf.get(term, 0.0) * (freq * 2.2) / (freq + 1.2 * (1 - 0.75 + 0.75 * doc_len / max(self.avgdl, 1)))
            if score > 0:
                hits.append(Hit(chunk, score))
        return sorted(hits, key=lambda hit: hit.score, reverse=True)[:limit]


def split_sentences(markdown: str) -> list[tuple[int, str]]:
    found: list[tuple[int, str]] = []
    in_fence = False
    in_references = False
    for line_no, line in enumerate(markdown.splitlines(), start=1):
        if line.strip().startswith("```"):
            in_fence = not in_fence
            continue
        stripped = re.sub(r"^\s*(?:[-*+] |\d+[.)]\s+|#+\s+|>\s*)", "", line).strip()
        if re.fullmatch(r"references?|참고문헌", stripped, re.I):
            in_references = True
            continue
        if in_references:
            continue
        if in_fence or not stripped or stripped.startswith("|"):
            continue
        if re.match(r"^(?:[①-⑳]|\(|x축:|y축:|색으로|각 방법의|환산 기준|실제 RGB)", stripped, re.I):
            continue
        for sentence in re.split(r"(?<=[.!?])\s+(?=[가-힣A-Z0-9])", stripped):
            sentence = normalize_space(sentence)
            # Exclude bare section titles such as "Ligand Crosslinking".
            if len(sentence) >= 8 and (re.search(r"[.!?]$", sentence) or len(sentence.split()) >= 6):
                found.append((line_no, sentence))
    return found


def claim_categories(sentence: str) -> list[str]:
    categories = [name for name, pattern in CLAIM_RULES.items() if pattern.search(sentence)]
    if categories and re.search(r"\[(?:@|\d)", sentence):
        categories.append("existing_citation")
    return categories


def citation_need_score(sentence: str, categories: list[str] | None = None) -> tuple[int, list[str]]:
    """Estimate whether a sentence makes an externally checkable claim.

    The score is intentionally conservative: statements about the paper's own
    organization or purpose are excluded, while numbers, named attributions,
    mechanisms, causal claims, and concrete technical limitations score higher.
    """
    categories = categories if categories is not None else claim_categories(sentence)
    score = 0
    reasons: list[str] = []
    weights = {
        "quantitative": (30, "정량 수치 포함"),
        "comparison": (20, "비교·성능 판단 포함"),
        "causal": (20, "인과관계 주장"),
        "generalization": (15, "일반화된 사실 주장"),
        "mechanism": (25, "작동 원리·메커니즘 설명"),
        "existing_citation": (30, "기존 인용 검증 필요"),
    }
    for category in categories:
        if category in weights:
            weight, reason = weights[category]
            score += weight
            reasons.append(reason)

    if re.search(r"\b(?:et al\.|demonstrated|reported|according to|as categorized by)\b", sentence, re.I):
        score += 20
        reasons.append("특정 선행연구 귀속")
    if re.search(r"\b(?:require(?:s|d)?|demand(?:s|ed)?|limited|limitation|drawback|suffer(?:s|ed)?)\b", sentence, re.I):
        score += 15
        reasons.append("요구조건·한계 주장")
    if re.search(r"\b(?:pixel|ppi|eqe|ligand|bandgap|resolution|luminance|photoluminescence|pattern(?:ing|ed)?)\b", sentence, re.I):
        score += 10
        reasons.append("구체적 기술 사실")
    if re.search(r"\bsignificant drawbacks\b", sentence, re.I) and not re.search(r"\b(?:because|due to|such as|including|while|whereas)\b", sentence, re.I):
        score -= 15
        reasons.append("근거 대상이 불명확한 포괄 표현")
    if re.search(r"^(?:The goal (?:is|of)|This review (?:examines|provides|focuses|organizes)|We (?:review|examine|organize|aim)|From an engineering standpoint, realizing .* across three interconnected dimensions)\b", sentence, re.I):
        score -= 60
        reasons.append("논문 자체의 목적·구성 설명")

    return max(0, min(score, 100)), reasons


def section_ranges(text: str) -> dict[str, tuple[int, int]]:
    """Return 1-based line ranges for the abstract and main body when detectable."""
    lines = text.splitlines()
    abstract_start = None
    body_heading = None
    body_start = None
    references_start = len(lines) + 1
    for number, line in enumerate(lines, start=1):
        stripped = normalize_space(line)
        if abstract_start is None and stripped.lower() == "abstract":
            # Ignore a table-of-contents entry if another Abstract appears later.
            later = any(normalize_space(item).lower() == "abstract" for item in lines[number:])
            if not later:
                abstract_start = number + 1
        if re.match(r"^(?:1\.?\s*)?introduction$", stripped, re.I):
            body_heading = number
            body_start = number + 1
        if re.fullmatch(r"references?|참고문헌", stripped, re.I):
            references_start = number
            break
    ranges: dict[str, tuple[int, int]] = {}
    if abstract_start and body_heading and abstract_start < body_heading:
        ranges["abstract"] = (abstract_start, body_heading - 1)
    if body_start:
        ranges["body"] = (body_start, references_start - 1)
    return ranges


def line_in_range(line: int, bounds: tuple[int, int] | None) -> bool:
    return bool(bounds and bounds[0] <= line <= bounds[1])


def extract_citation_markers(sentence: str) -> list[str]:
    return re.findall(r"\[(?:\d+(?:\s*[-,]\s*\d+)*)\]", sentence)


def extract_citation_numbers(sentence: str) -> list[int]:
    numbers: list[int] = []
    for marker in extract_citation_markers(sentence):
        content = marker[1:-1]
        for part in re.split(r"\s*,\s*", content):
            if "-" in part:
                start, end = (int(value.strip()) for value in part.split("-", 1))
                numbers.extend(range(start, end + 1))
            elif part.strip().isdigit():
                numbers.append(int(part.strip()))
    return list(dict.fromkeys(numbers))


def parse_numbered_references(text: str) -> dict[int, str]:
    references: dict[int, str] = {}
    in_references = False
    current: int | None = None
    for line in text.splitlines():
        stripped = normalize_space(line)
        if re.fullmatch(r"references?|참고문헌", stripped, re.I):
            in_references = True
            continue
        if not in_references or not stripped:
            continue
        match = re.match(r"^\[(\d+)\]\s*(.+)$", stripped)
        if match:
            current = int(match.group(1))
            references[current] = match.group(2)
        elif current is not None:
            references[current] += " " + stripped
    return references


def match_reference_pdf(reference: str, engine: BM25) -> tuple[str, float] | None:
    reference_doi = re.search(r"10\.\d{4,9}/[-._;()/:A-Z0-9]+", reference, re.I)
    reference_terms = set(tokenize(reference))
    by_file: dict[str, Chunk] = {}
    for chunk in engine.chunks:
        by_file.setdefault(chunk.file, chunk)

    scored: list[tuple[str, float]] = []
    for filename, chunk in by_file.items():
        if reference_doi and chunk.doi and reference_doi.group(0).rstrip(".,;").lower() == chunk.doi.rstrip(".,;").lower():
            return filename, 100.0
        candidate = f"{Path(filename).stem} {chunk.title}"
        candidate_terms = set(tokenize(candidate))
        overlap = len(reference_terms & candidate_terms) / max(min(len(reference_terms), len(candidate_terms)), 1)
        sequence = SequenceMatcher(None, normalize_space(reference).lower(), normalize_space(candidate).lower()).ratio()
        scored.append((filename, 70 * overlap + 30 * sequence))

    if not scored:
        return None
    best = max(scored, key=lambda item: item[1])
    # A weak best guess is more misleading than an explicit missing-PDF status.
    return best if best[1] >= 38 else None


def search_within_pdf(claim: str, filename: str, chunks: list[Chunk], limit: int = 2) -> list[Hit]:
    selected = [chunk for chunk in chunks if chunk.file == filename]
    return BM25(selected).search(claim, limit) if selected else []


def claim_key(claim: str) -> str:
    return hashlib.sha1(normalize_space(claim).encode("utf-8")).hexdigest()[:12]


def verdict_file_for(draft: Path) -> Path:
    return OUTPUT / f"{draft.stem}_ai_verdicts.json"


def load_ai_verdicts(draft: Path) -> dict[str, dict]:
    path = verdict_file_for(draft)
    if not path.exists():
        return {}
    data = json.loads(path.read_text(encoding="utf-8"))
    return {item["claim_id"]: item for item in data.get("verdicts", [])}


def prepare_ai_review(draft: Path, limit: int = 3) -> Path:
    text = load_draft_text(draft)
    ranges = section_ranges(text)
    claims = [(line, sentence, claim_categories(sentence)) for line, sentence in split_sentences(text)]
    claims = [
        (line, sentence, cats)
        for line, sentence, cats in claims
        if cats
        and citation_need_score(sentence, cats)[0] >= CITATION_NEED_THRESHOLD
        and not line_in_range(line, ranges.get("abstract"))
    ]
    references = parse_numbered_references(text)
    engine = BM25(load_index())
    items = []
    for line, claim, categories in claims:
        cited = []
        for number in extract_citation_numbers(claim):
            reference = references.get(number, "")
            match = match_reference_pdf(reference, engine) if reference else None
            evidence = []
            if match:
                evidence = [
                    {"file": hit.chunk.file, "page": hit.chunk.page, "excerpt": relevant_excerpt(claim, hit.chunk.text, 900)}
                    for hit in search_within_pdf(claim, match[0], engine.chunks, 3)
                ]
            cited.append({"number": number, "reference": reference, "matched_pdf": match[0] if match else "", "evidence": evidence})
        alternatives = [
            {"file": hit.chunk.file, "page": hit.chunk.page, "excerpt": relevant_excerpt(claim, hit.chunk.text, 700)}
            for hit in engine.search(claim, limit)
        ]
        items.append({
            "claim_id": claim_key(claim),
            "line": line,
            "claim": claim,
            "categories": categories,
            "citation_need": {
                "score": citation_need_score(claim, categories)[0],
                "threshold": CITATION_NEED_THRESHOLD,
                "reasons": citation_need_score(claim, categories)[1],
            },
            "existing_citations": cited,
            "alternative_candidates": alternatives,
            "required_output": {
                "verdict": "direct|partial|mismatch|insufficient",
                "supported_parts": ["..."],
                "unsupported_parts": ["..."],
                "reason": "...",
                "recommended_action": "keep|revise|split|add_reference|replace_reference|external_search",
                "suggested_revision": "...",
                "confidence": "high|medium|low",
                "followup_search": {
                    "status": "not_needed|pending|completed",
                    "searched_gaps": ["..."],
                    "result": "evidence_found|no_suitable_evidence",
                    "summary": "..."
                },
                "external_search": {
                    "status": "not_needed|pending|completed",
                    "proposals": [{
                        "citation": "...",
                        "url": "DOI or publisher page",
                        "full_text_url": "lawful full-text URL or empty",
                        "full_text_status": "acquired_verified|available_not_acquired|not_found",
                        "local_file": "filename in papers/ or empty",
                        "verification": "title/author/DOI and body-evidence check",
                        "supports": "...",
                        "source_location": "page/section",
                        "source_excerpt": "...",
                        "limitation": "...",
                        "proposed_revision": "...",
                        "renumbering_plan": "..."
                    }]
                }
            },
        })
    destination = OUTPUT / f"{draft.stem}_ai_review_packet.json"
    destination.write_text(json.dumps({"draft": str(draft.resolve()), "items": items}, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Codex 검토 패킷: {destination}")
    return destination


def body_support_hits(text: str, abstract_claim: str, limit: int = 3) -> list[tuple[int, str, list[str], float]]:
    ranges = section_ranges(text)
    body = ranges.get("body")
    candidates = [(line, sentence) for line, sentence in split_sentences(text) if line_in_range(line, body)]
    chunks = [Chunk(str(index), "body", line, sentence) for index, (line, sentence) in enumerate(candidates)]
    if not chunks:
        return []
    hits = BM25(chunks).search(abstract_claim, limit)
    return [(hit.chunk.page, hit.chunk.text, extract_citation_markers(hit.chunk.text), hit.score) for hit in hits]


def excerpt(text: str, width: int = 360) -> str:
    clean = normalize_space(text)
    return clean if len(clean) <= width else clean[: width - 3] + "..."


def relevant_excerpt(query: str, text: str, width: int = 520) -> str:
    """Extract the sentence/window most relevant to a claim and remove boilerplate."""
    clean = normalize_space(text)
    clean = re.sub(r"https?://\S+|www\.\S+", " ", clean, flags=re.I)
    clean = re.sub(r"\bdoi\s*:?\s*10\.\d{4,9}/\S+", " ", clean, flags=re.I)
    clean = re.sub(r"\b(?:official journal|open access|read online|article recommendations|supporting information|metrics\s*&\s*more)\b", " ", clean, flags=re.I)
    clean = normalize_space(clean)
    abstract_match = re.search(r"\bAbstract\b", clean[:500], re.I)
    if abstract_match:
        clean = clean[abstract_match.end() :].strip()
    query_terms = set(tokenize(query))
    candidates = [normalize_space(item) for item in re.split(r"(?<=[.!?])\s+", clean) if len(normalize_space(item)) >= 25]
    words = clean.split()
    if len(words) > 55:
        candidates.extend(" ".join(words[start : start + 55]) for start in range(0, len(words), 28))
    if not candidates:
        return excerpt(clean, width)

    def score(candidate: str) -> tuple[int, float]:
        terms = set(tokenize(candidate))
        overlap = len(query_terms & terms)
        return overlap, overlap / max(len(terms), 1)

    return excerpt(max(candidates, key=score), width)


def append_audit(row: dict[str, str]) -> None:
    exists = AUDIT_FILE.exists()
    fields = ["timestamp", "draft", "line", "claim", "categories", "decision", "source_file", "source_page", "source_doi", "revised_claim", "note"]
    with AUDIT_FILE.open("a", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields)
        if not exists:
            writer.writeheader()
        writer.writerow({field: row.get(field, "") for field in fields})


def enqueue_web_search(draft: Path, line: int, claim: str, categories: list[str], note: str) -> None:
    request = {
        "created_at": datetime.now(timezone.utc).isoformat(),
        "draft": str(draft.resolve()),
        "line": line,
        "claim": claim,
        "categories": categories,
        "user_note": note,
        "requirements": [
            "원저 논문 우선",
            "제목·저자·연도·DOI 확인",
            "주장을 뒷받침하는 본문·페이지·표·Figure 위치 확인",
            "검색결과 요약만으로 reference 확정 금지",
        ],
        "status": "pending",
    }
    with WEB_QUEUE.open("a", encoding="utf-8") as handle:
        handle.write(json.dumps(request, ensure_ascii=False) + "\n")


def replace_once(text: str, old: str, new: str) -> str:
    position = text.find(old)
    if position < 0:
        return text
    return text[:position] + new + text[position + len(old):]


def load_draft_text(draft: Path) -> str:
    if draft.suffix.lower() in {".md", ".txt"}:
        return draft.read_text(encoding="utf-8")
    if draft.suffix.lower() == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise SystemExit("python-docx가 필요함: python -m pip install python-docx") from exc
        document = Document(draft)
        blocks = [paragraph.text for paragraph in iter_docx_paragraphs(document)]
        return "\n".join(blocks)
    raise SystemExit("지원하지 않는 초안 형식임. .docx, .md, .txt를 사용해줘.")


def iter_docx_paragraphs(document) -> Iterable:
    """Yield body and table-cell paragraphs in their practical reading order."""
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def replace_paragraph_text(paragraph, old: str, new: str) -> bool:
    full = paragraph.text
    if old not in full:
        return False
    replaced = full.replace(old, new, 1)
    # Only changed paragraphs are rebuilt; the paragraph style and layout remain.
    if paragraph.runs:
        paragraph.runs[0].text = replaced
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(replaced)
    return True


def save_docx_review(source: Path, destination: Path, replacements: list[tuple[str, str]]) -> None:
    from docx import Document

    shutil.copy2(source, destination)
    document = Document(destination)
    paragraphs = list(iter_docx_paragraphs(document))
    for old, new in replacements:
        for paragraph in paragraphs:
            if replace_paragraph_text(paragraph, old, new):
                break
    document.save(destination)


def review(draft: Path) -> None:
    ensure_dirs()
    if not draft.exists():
        raise SystemExit(f"초안 파일이 없음: {draft}")
    original = load_draft_text(draft)
    reviewed = original
    replacements: list[tuple[str, str]] = []
    ranges = section_ranges(original)
    claims = [(line, sentence, claim_categories(sentence)) for line, sentence in split_sentences(original)]
    claims = [
        (line, sentence, categories)
        for line, sentence, categories in claims
        if categories
        and citation_need_score(sentence, categories)[0] >= CITATION_NEED_THRESHOLD
        and not line_in_range(line, ranges.get("abstract"))
    ]
    if not claims:
        print("인용 필요 후보를 찾지 못함. 규칙을 검토하거나 문장에 [REFERENCE REQUIRED]를 직접 표시해줘.")
        return

    engine = BM25(load_index())
    output_suffix = ".docx" if draft.suffix.lower() == ".docx" else ".md"
    output_path = OUTPUT / f"{draft.stem}_reviewed{output_suffix}"
    print(f"인용 필요 후보 {len(claims)}개")

    for number, (line, claim, categories) in enumerate(claims, start=1):
        hits = engine.search(claim, 3)
        print("\n" + "=" * 78)
        print(f"[{number}/{len(claims)}] line {line} | {', '.join(categories)}")
        print(claim)
        if hits:
            for idx, hit in enumerate(hits, start=1):
                meta = f"DOI {hit.chunk.doi}" if hit.chunk.doi else "DOI 미추출"
                print(f"\n{idx}. {hit.chunk.file}, p.{hit.chunk.page}, score={hit.score:.2f}, {meta}")
                print(f"   {excerpt(hit.chunk.text)}")
        else:
            print("\n내부 근거 후보 없음")

        allowed = {str(index) for index in range(1, len(hits) + 1)} | {"r", "w", "u", "s", "q"}
        while True:
            choice = input("\n선택 [1-3 근거채택 / r 수정 / w 외부검색 / u 미검증 / s 제외 / q 종료]: ").strip().lower()
            if choice in allowed:
                break
            print("유효한 선택이 아님")

        common = {
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "draft": str(draft.resolve()),
            "line": str(line),
            "claim": claim,
            "categories": ";".join(categories),
        }
        if choice == "q":
            if draft.suffix.lower() == ".docx":
                save_docx_review(draft, output_path, replacements)
            else:
                output_path.write_text(reviewed, encoding="utf-8")
            print(f"중간 저장: {output_path}")
            return
        if choice.isdigit():
            hit = hits[int(choice) - 1]
            key = f"[@{Path(hit.chunk.file).stem}-p{hit.chunk.page}]"
            replacement = f"{claim} {key}"
            reviewed = replace_once(reviewed, claim, replacement)
            replacements.append((claim, replacement))
            append_audit({**common, "decision": "accepted_internal_evidence", "source_file": hit.chunk.file, "source_page": str(hit.chunk.page), "source_doi": hit.chunk.doi})
        elif choice == "r":
            revised = input("수정 문장: ").strip()
            if revised:
                reviewed = replace_once(reviewed, claim, revised)
                replacements.append((claim, revised))
            append_audit({**common, "decision": "revised_by_user", "revised_claim": revised})
        elif choice == "w":
            note = input("외부검색 조건(선택, Enter로 생략): ").strip()
            enqueue_web_search(draft, line, claim, categories, note)
            replacement = f"{claim} [EXTERNAL SEARCH PENDING]"
            reviewed = replace_once(reviewed, claim, replacement)
            replacements.append((claim, replacement))
            append_audit({**common, "decision": "external_search_requested", "note": note})
        elif choice == "u":
            replacement = f"{claim} [REFERENCE REQUIRED]"
            reviewed = replace_once(reviewed, claim, replacement)
            replacements.append((claim, replacement))
            append_audit({**common, "decision": "kept_unverified"})
        else:
            append_audit({**common, "decision": "citation_not_needed"})

        if draft.suffix.lower() == ".docx":
            save_docx_review(draft, output_path, replacements)
        else:
            output_path.write_text(reviewed, encoding="utf-8")

    print(f"\n검토 완료: {output_path}")
    print(f"감사 로그: {AUDIT_FILE}")
    if WEB_QUEUE.exists():
        print(f"외부검색 요청: {WEB_QUEUE}")


def scan(draft: Path, limit: int = 3) -> Path:
    ensure_dirs()
    text = load_draft_text(draft)
    ranges = section_ranges(text)
    all_candidates = [(line, sentence, claim_categories(sentence)) for line, sentence in split_sentences(text)]
    all_candidates = [
        item for item in all_candidates
        if item[2] and not line_in_range(item[0], ranges.get("abstract"))
    ]
    claims = all_candidates
    claims = [
        (line, sentence, categories)
        for line, sentence, categories in claims
        if categories
        and citation_need_score(sentence, categories)[0] >= CITATION_NEED_THRESHOLD
        and not line_in_range(line, ranges.get("abstract"))
    ]
    engine = BM25(load_index())
    ai_verdicts = load_ai_verdicts(draft)
    completed_count = sum(1 for _, claim, _ in claims if claim_key(claim) in ai_verdicts)
    report = [
        f"# {draft.name} 인용 검토",
        "",
        f"- 검토 대상: {len(claims)}개",
        f"- 보수적 필터 제외: {len(all_candidates) - len(claims)}개",
        f"- 포함 기준: 인용 필요도 {CITATION_NEED_THRESHOLD}/100 이상",
        f"- 판정 완료: {completed_count}개",
        f"- 판정 대기: {max(len(claims) - completed_count, 0)}개",
        "",
    ]
    for number, (line, claim, categories) in enumerate(claims, start=1):
        need_score, need_reasons = citation_need_score(claim, categories)
        existing_markers = extract_citation_markers(claim)
        existing_numbers = extract_citation_numbers(claim)
        numbered_references = parse_numbered_references(text)
        existing_text = ", ".join(existing_markers) if existing_markers else "없음"
        verdict = ai_verdicts.get(claim_key(claim))
        report.extend([
            f"## {number}",
            "",
            f"> {claim}",
            "",
            f"- 기존 인용: `{existing_text}`",
            f"- 인용 필요도: `{need_score}/100` (검토 기준 {CITATION_NEED_THRESHOLD}점)",
            f"- 포함 이유: {', '.join(need_reasons)}",
        ])
        if not verdict:
            report.extend(["- 판정: `대기`", ""])
            continue

        report.extend([
            f"- 판정: `{display_label(verdict.get('verdict', ''), VERDICT_LABELS)}`",
            f"- 판단 이유: {verdict.get('reason', '')}",
            f"- 뒷받침되는 부분: {'; '.join(verdict.get('supported_parts', [])) or '없음'}",
            f"- 부족한 부분: {'; '.join(verdict.get('unsupported_parts', [])) or '없음'}",
        ])
        local_evidence = verdict.get("local_evidence") or []
        if local_evidence:
            report.extend(["", "### 확인된 근거", ""])
            for evidence in local_evidence:
                report.append(
                    f"- {evidence.get('file', '')}, p.{evidence.get('page', '')}: “{evidence.get('excerpt', '')}”"
                )
        followup = verdict.get("followup_search") or {}
        if verdict.get("verdict") == "partial":
            followup_status = followup.get("status", "pending")
            status_label = {"pending": "대기", "completed": "완료", "not_needed": "불필요"}.get(followup_status, followup_status)
            report.append(f"- 부족한 근거 재검색: `{status_label}`")
            if followup.get("summary"):
                report.append(f"- 재검색 결과: {followup['summary']}")
        if existing_numbers:
            for citation_number in existing_numbers:
                reference = numbered_references.get(citation_number)
                if not reference:
                    report.append(f"- [{citation_number}] 근거: 참고문헌 항목을 찾지 못함")
                    continue
                match = match_reference_pdf(reference, engine)
                if not match:
                    report.append(f"- [{citation_number}] 근거: 대응 PDF를 찾지 못함")
                    continue
                matched_file, match_score = match
                within_hits = search_within_pdf(claim, matched_file, engine.chunks, 1)
                if not within_hits:
                    report.append(f"- [{citation_number}] 근거 구절: 찾지 못함")
                for hit in within_hits:
                    report.append(f"- [{citation_number}] 근거 구절: “{relevant_excerpt(claim, hit.chunk.text, 520)}” ({Path(matched_file).stem}, p.{hit.chunk.page})")
        external = verdict.get("external_search") or {}
        if external.get("status") == "completed":
            report.extend(["", "### 외부 논문 제안", ""])
            proposals = external.get("proposals", [])
            if not proposals:
                report.append("- 적합한 외부 논문을 찾지 못함")
            for proposal_number, proposal in enumerate(proposals, start=1):
                full_text_labels = {
                    "acquired_verified": "원문 확보·본문 확인 완료",
                    "available_not_acquired": "합법적 원문 경로 확인·미확보",
                    "not_found": "원문 미확보",
                }
                full_text_status = proposal.get("full_text_status", "not_found")
                report.extend([
                    f"#### 제안 {proposal_number}",
                    "",
                    f"- 논문: {proposal.get('citation', '')}",
                    f"- 링크: {proposal.get('url', '')}",
                    f"- 원문 상태: `{full_text_labels.get(full_text_status, full_text_status)}`",
                    f"- 합법적 원문: {proposal.get('full_text_url', '') or '찾지 못함'}",
                    f"- 저장 파일: {proposal.get('local_file', '') or '없음'}",
                    f"- 원문 검증: {proposal.get('verification', '') or '미검증'}",
                    f"- 뒷받침하는 내용: {proposal.get('supports', '')}",
                    f"- 원문 위치: {proposal.get('source_location', '')}",
                    f"- 원문 구절: “{proposal.get('source_excerpt', '')}”",
                    f"- 한계: {proposal.get('limitation', '')}",
                    f"- 적용 문장: {proposal.get('proposed_revision', '')}",
                    f"- 번호 변경안: {proposal.get('renumbering_plan', '')}",
                    "- 결정: `대기` (`채택 / 수정 후 채택 / 거절`)",
                    "",
                ])
        report.extend([
            f"- 권장 조치: `{display_label(verdict.get('recommended_action', ''), ACTION_LABELS)}`",
            f"- 수정 제안: {verdict.get('suggested_revision', '') or '없음'}",
            "- 사용자 결정: `대기`",
            "",
        ])
    destination = OUTPUT / f"{draft.stem}_evidence_scan.md"
    destination.write_text("\n".join(report), encoding="utf-8")
    print(f"사전 점검 보고서: {destination}")
    return destination


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="보유 논문 기반 졸업논문 근거 추적 도우미")
    sub = parser.add_subparsers(dest="command", required=True)
    def add_paths(command_parser: argparse.ArgumentParser) -> None:
        command_parser.add_argument("--papers-dir", type=Path, help="레퍼런스 PDF 폴더")
        command_parser.add_argument("--output-dir", type=Path, help="색인·보고서 저장 폴더")

    init_parser = sub.add_parser("init", help="작업 폴더 생성")
    add_paths(init_parser)
    index_parser = sub.add_parser("index", help="papers 폴더의 PDF 색인")
    add_paths(index_parser)
    review_parser = sub.add_parser("review", help="Word/Markdown/텍스트 초안 대화형 검토")
    add_paths(review_parser)
    review_parser.add_argument("draft", type=Path)
    scan_parser = sub.add_parser("scan", help="초안 전체의 근거 후보 보고서 생성")
    add_paths(scan_parser)
    scan_parser.add_argument("draft", type=Path)
    scan_parser.add_argument("--limit", type=int, default=3)
    prepare_parser = sub.add_parser("prepare-ai", help="Codex 의미 검토용 JSON 패킷 생성")
    add_paths(prepare_parser)
    prepare_parser.add_argument("draft", type=Path)
    prepare_parser.add_argument("--limit", type=int, default=3)
    search_parser = sub.add_parser("search", help="내부 근거 검색 테스트")
    add_paths(search_parser)
    search_parser.add_argument("query")
    return parser


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    configure_paths(getattr(args, "papers_dir", None), getattr(args, "output_dir", None))
    ensure_dirs()
    if args.command == "init":
        print(f"생성/확인 완료: {PAPERS}, {DRAFTS}, {OUTPUT}")
    elif args.command == "index":
        index_papers()
    elif args.command == "review":
        review(args.draft)
    elif args.command == "scan":
        scan(args.draft, args.limit)
    elif args.command == "prepare-ai":
        prepare_ai_review(args.draft, args.limit)
    elif args.command == "search":
        engine = BM25(load_index())
        for idx, hit in enumerate(engine.search(args.query, 5), start=1):
            print(f"{idx}. {hit.chunk.file} p.{hit.chunk.page} score={hit.score:.2f}")
            print(excerpt(hit.chunk.text))
    return 0


if __name__ == "__main__":
    sys.exit(main())
