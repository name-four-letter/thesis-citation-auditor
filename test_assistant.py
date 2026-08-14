import unittest
from pathlib import Path

from thesis_assistant import (
    BM25,
    Chunk,
    CITATION_NEED_THRESHOLD,
    citation_need_score,
    claim_categories,
    load_draft_text,
    save_docx_review,
    split_chunks,
    split_sentences,
)
from thesis_assistant import body_support_hits, extract_citation_numbers, parse_numbered_references, relevant_excerpt, section_ranges


class ThesisAssistantTests(unittest.TestCase):
    def test_claim_detection(self):
        categories = claim_categories("선행연구에서는 최소 2 μm의 패턴 폭이 보고되었다.")
        self.assertIn("quantitative", categories)
        self.assertIn("generalization", categories)

    def test_existing_citation_is_marked(self):
        self.assertIn("existing_citation", claim_categories("최소 패턴 폭은 2 μm이다 [@kim2024]."))

    def test_authorial_scope_statement_is_below_threshold(self):
        score, reasons = citation_need_score(
            "The goal is to provide a comprehensive overview of the current landscape."
        )
        self.assertLess(score, CITATION_NEED_THRESHOLD)
        self.assertIn("논문 자체의 목적·구성 설명", reasons)

    def test_quantitative_technical_claim_is_reviewed(self):
        sentence = "Near-eye displays require pixel densities exceeding 5,000 PPI [2]."
        score, _ = citation_need_score(sentence)
        self.assertGreaterEqual(score, CITATION_NEED_THRESHOLD)

    def test_references_section_is_skipped(self):
        text = "본문은 5,000 PPI가 필요하다고 보고되었다.\nReferences\n[1] Reported 10,000 PPI display."
        sentences = split_sentences(text)
        self.assertEqual(len(sentences), 1)

    def test_abstract_maps_to_cited_body_claim(self):
        text = (
            "Abstract\n"
            "Near-eye displays require more than 5,000 PPI.\n"
            "1. Introduction\n"
            "AR and VR microdisplays require pixel densities exceeding 5,000 PPI [2].\n"
            "References\n[2] Example reference."
        )
        self.assertEqual(section_ranges(text)["abstract"], (2, 2))
        hits = body_support_hits(text, "Near-eye displays require more than 5,000 PPI.")
        self.assertEqual(hits[0][2], ["[2]"])

    def test_citation_ranges_and_reference_parsing(self):
        self.assertEqual(extract_citation_numbers("성과가 보고되었다 [2, 5-7]."), [2, 5, 6, 7])
        text = "본문 [1].\nReferences\n[1] First paper title.\ncontinued metadata\n[2] Second paper."
        references = parse_numbered_references(text)
        self.assertIn("continued metadata", references[1])
        self.assertEqual(references[2], "Second paper.")

    def test_relevant_excerpt_removes_boilerplate(self):
        text = (
            "Official journal Open Access https://doi.org/10.1000/example. "
            "Unrelated publishing information appears here. "
            "AR and VR headsets are emerging as next-generation interactive displays capable of vivid 3D experiences."
        )
        result = relevant_excerpt("AR VR interactive displays vivid 3D experiences", text)
        self.assertIn("interactive displays", result)
        self.assertNotIn("https://", result)
        self.assertNotIn("Official journal", result)

    def test_sentence_split_tracks_line(self):
        sentences = split_sentences("# 제목\n\n해당 공정은 기존 방식보다 효율적이다. 결과가 향상되었다.")
        self.assertEqual(sentences[0][0], 3)
        self.assertGreaterEqual(len(sentences), 2)

    def test_chunk_size(self):
        text = "첫 번째 문단은 충분히 긴 기술 설명을 포함하고 있다. " * 20
        chunks = split_chunks(text, max_chars=200)
        self.assertGreater(len(chunks), 1)

    def test_bm25_prefers_relevant_evidence(self):
        chunks = [
            Chunk("1", "a.pdf", 1, "Direct optical lithography achieved a minimum pattern width of 2 μm."),
            Chunk("2", "b.pdf", 2, "Quantum dot emission wavelength and color purity were measured."),
        ]
        hit = BM25(chunks).search("최소 pattern width 2 μm", 1)[0]
        self.assertEqual(hit.chunk.id, "1")

    def test_docx_round_trip_preserves_unmodified_paragraph(self):
        from docx import Document

        test_tmp_root = Path(__file__).resolve().parent / "output"
        test_tmp_root.mkdir(exist_ok=True)
        source = test_tmp_root / "_test_source.docx"
        output = test_tmp_root / "_test_reviewed.docx"
        try:
            document = Document()
            document.add_heading("검토 예시", level=1)
            document.add_paragraph("최소 패턴 폭은 2 μm이다.")
            document.add_paragraph("수정하지 않는 문단이다.")
            document.save(source)

            extracted = load_draft_text(source)
            self.assertIn("최소 패턴 폭은 2 μm이다.", extracted)
            save_docx_review(source, output, [("최소 패턴 폭은 2 μm이다.", "최소 패턴 폭은 2 μm이다. [REFERENCE REQUIRED]")])
            reviewed = Document(output)
            self.assertEqual(reviewed.paragraphs[2].text, "수정하지 않는 문단이다.")
            self.assertIn("[REFERENCE REQUIRED]", reviewed.paragraphs[1].text)
        finally:
            source.unlink(missing_ok=True)
            output.unlink(missing_ok=True)


if __name__ == "__main__":
    unittest.main()
